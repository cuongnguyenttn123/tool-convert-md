#!/usr/bin/env python3
"""Convert DOCX, Excel, or PDF files to Markdown.

Rules:
- DOCX: output a folder with the same base name, containing a markdown file
  of the same name plus an images/ folder. Supports headings, lists, tables
  (including nested tables, rendered as HTML), text boxes, headers/footers,
  footnotes/endnotes, hyperlinks (including internal bookmarks), and images.
- Excel: output a folder with the same base name. Each sheet becomes one
  markdown file, with a shared images/ folder for extracted images.
- PDF: output a folder with the same base name, containing a markdown file
  of the same name plus an images/ folder for extracted images.
"""

from __future__ import annotations

import argparse
import html as html_mod
import re
import shutil
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Iterable, Optional

import pymupdf4llm
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree
from openpyxl import load_workbook

# --------------------------------------------------------------------------
# Shared helpers
# --------------------------------------------------------------------------

INLINE_CONTAINER_TAGS = {qn("w:ins"), qn("w:smartTag"), qn("w:sdt"), qn("w:sdtContent")}
NOTE_SKIP_TYPES = {"separator", "continuationSeparator", "continuationNotice"}


def sanitize_name(name: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|]", "_", name.strip())
    return name or "untitled"


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def reset_dir(path: Path) -> None:
    """(Re)create an empty directory, so a re-run never leaves stale files
    (e.g. images removed from a since-edited source document) behind."""
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True)


def write_text(path: Path, content: str) -> None:
    path.write_text(content.rstrip() + "\n", encoding="utf-8")


@dataclass
class ConvertState:
    image_dir: Path
    image_index: int = 0
    used_footnotes: list[str] = field(default_factory=list)
    used_endnotes: list[str] = field(default_factory=list)


@dataclass
class InlineResolver:
    resolve_link: Callable[[object], Optional[str]]
    get_image_part: Callable[[str], Optional[object]] = lambda rel_id: None


def _part_resolver(part) -> InlineResolver:
    def resolve_link(hyperlink_elem) -> Optional[str]:
        r_id = hyperlink_elem.get(qn("r:id"))
        if r_id:
            try:
                return part.rels[r_id].target_ref
            except KeyError:
                return None
        anchor = hyperlink_elem.get(qn("w:anchor"))
        return f"#{anchor}" if anchor else None

    def get_image_part(rel_id: str):
        return part.related_parts.get(rel_id)

    return InlineResolver(resolve_link=resolve_link, get_image_part=get_image_part)


def _rels_resolver(rels: dict[str, str]) -> InlineResolver:
    def resolve_link(hyperlink_elem) -> Optional[str]:
        r_id = hyperlink_elem.get(qn("r:id"))
        if r_id and r_id in rels:
            return rels[r_id]
        anchor = hyperlink_elem.get(qn("w:anchor"))
        return f"#{anchor}" if anchor else None

    return InlineResolver(resolve_link=resolve_link)


# --------------------------------------------------------------------------
# Inline content (text, hyperlinks, images, footnote/endnote refs)
# --------------------------------------------------------------------------


def _images_in(elem) -> list[tuple[str, str, None]]:
    segments = []
    for blip in elem.xpath(".//a:blip"):
        rel_id = blip.get(qn("r:embed"))
        if rel_id:
            segments.append(("image", rel_id, None))
    return segments


def _walk_inline(elem, resolver: InlineResolver) -> list[tuple[str, str, Optional[str]]]:
    """Walk one paragraph-like element's direct children into ordered segments."""
    segments: list[tuple[str, str, Optional[str]]] = []

    for child in elem.iterchildren():
        tag = child.tag

        if tag == qn("w:hyperlink"):
            text = "".join(child.xpath(".//w:t/text()"))
            if text:
                segments.append(("link", text, resolver.resolve_link(child)))
            segments.extend(_images_in(child))

        elif tag == qn("w:r"):
            fn_ref = child.find(qn("w:footnoteReference"))
            if fn_ref is not None:
                segments.append(("footnote", fn_ref.get(qn("w:id")), None))
            en_ref = child.find(qn("w:endnoteReference"))
            if en_ref is not None:
                segments.append(("endnote", en_ref.get(qn("w:id")), None))

            text = "".join(t.text or "" for t in child.findall(qn("w:t")))
            if child.find(qn("w:tab")) is not None:
                text += "\t"
            if text:
                segments.append(("text", text, None))
            if child.find(qn("w:br")) is not None or child.find(qn("w:cr")) is not None:
                segments.append(("break", "", None))

            segments.extend(_images_in(child))

        elif tag in INLINE_CONTAINER_TAGS:
            segments.extend(_walk_inline(child, resolver))

        elif tag == qn("w:del"):
            continue

    return segments


def _write_image(rel_id: str, resolver: InlineResolver, state: ConvertState) -> Optional[str]:
    img_part = resolver.get_image_part(rel_id)
    if img_part is None:
        return None
    ext = Path(img_part.filename).suffix or ".png"
    state.image_index += 1
    filename = f"img_{state.image_index:04d}{ext}"
    target = state.image_dir / filename
    target.write_bytes(img_part.blob)
    return f"{state.image_dir.name}/{filename}"


def segments_to_markdown(segments, resolver: InlineResolver, state: ConvertState) -> str:
    out: list[str] = []
    for kind, text, url in segments:
        if kind == "break":
            out.append("  \n")
        elif kind == "footnote":
            if text not in state.used_footnotes:
                state.used_footnotes.append(text)
            out.append(f"[^fn{text}]")
        elif kind == "endnote":
            if text not in state.used_endnotes:
                state.used_endnotes.append(text)
            out.append(f"[^en{text}]")
        elif kind == "link":
            out.append(f"[{text}]({url})" if url else text)
        elif kind == "image":
            rel_path = _write_image(text, resolver, state)
            if rel_path:
                out.append(f"![image]({rel_path})")
        else:
            out.append(text)
    return "".join(out)


def _segments_to_html(segments, resolver: InlineResolver, state: ConvertState) -> str:
    out: list[str] = []
    for kind, text, url in segments:
        if kind == "break":
            out.append("<br>")
        elif kind == "footnote":
            if text not in state.used_footnotes:
                state.used_footnotes.append(text)
            out.append(f"[^fn{text}]")
        elif kind == "endnote":
            if text not in state.used_endnotes:
                state.used_endnotes.append(text)
            out.append(f"[^en{text}]")
        elif kind == "link":
            esc = html_mod.escape(text)
            if url:
                out.append(f'<a href="{html_mod.escape(url, quote=True)}">{esc}</a>')
            else:
                out.append(esc)
        elif kind == "image":
            rel_path = _write_image(text, resolver, state)
            if rel_path:
                out.append(f'<img src="{rel_path}" alt="image">')
        else:
            out.append(html_mod.escape(text))
    return "".join(out)


def paragraph_inline_markdown(paragraph: Paragraph, resolver: InlineResolver, state: ConvertState) -> str:
    return segments_to_markdown(_walk_inline(paragraph._p, resolver), resolver, state)


def _paragraph_inline_html(paragraph: Paragraph, resolver: InlineResolver, state: ConvertState) -> str:
    return _segments_to_html(_walk_inline(paragraph._p, resolver), resolver, state)


# --------------------------------------------------------------------------
# Text boxes (VML and DrawingML) embedded in a paragraph
# --------------------------------------------------------------------------


def _textbox_blocks(paragraph: Paragraph, resolver: InlineResolver, state: ConvertState) -> list[str]:
    blocks: list[str] = []
    for txbx in paragraph._p.xpath(".//*[local-name()='txbxContent']"):
        # mc:AlternateContent duplicates textbox content in Choice (DrawingML)
        # and Fallback (VML); only render the Fallback copy when no Choice sibling exists.
        if txbx.xpath("ancestor::*[local-name()='Fallback']"):
            continue

        lines: list[str] = []
        for child in txbx.iterchildren():
            if child.tag == qn("w:p"):
                text = segments_to_markdown(_walk_inline(child, resolver), resolver, state).strip()
                if text:
                    lines.append(text)
            elif child.tag == qn("w:tbl"):
                for tr in child.findall(qn("w:tr")):
                    cells = tr.findall(qn("w:tc"))
                    cell_texts = ["".join(tc.xpath(".//w:t/text()")).strip() for tc in cells]
                    lines.append(" | ".join(cell_texts))

        if lines:
            quoted = "\n".join(f"> {line}" if line else ">" for line in lines)
            blocks.append(f"> **[Text Box]**\n>\n{quoted}")

    return blocks


# --------------------------------------------------------------------------
# Paragraphs and tables (with nested-table support)
# --------------------------------------------------------------------------


def paragraph_to_markdown(paragraph: Paragraph, resolver: InlineResolver, state: ConvertState) -> list[str]:
    style_name = paragraph.style.name if paragraph.style else ""
    text = paragraph_inline_markdown(paragraph, resolver, state).strip()
    lines: list[str] = []

    if text:
        heading_match = re.match(r"Heading\s+(\d+)$", style_name)
        if heading_match:
            level = max(1, min(6, int(heading_match.group(1))))
            lines.append(f"{'#' * level} {text}")
        elif "List Bullet" in style_name:
            lines.append(f"- {text}")
        elif "List Number" in style_name:
            lines.append(f"1. {text}")
        else:
            lines.append(text)

    lines.extend(_textbox_blocks(paragraph, resolver, state))
    return lines


def _cell_grid_span(cell) -> int:
    tc_pr = cell._tc.tcPr
    if tc_pr is not None and tc_pr.gridSpan is not None:
        return tc_pr.gridSpan.val or 1
    return 1


def _table_to_html(table: Table, resolver: InlineResolver, state: ConvertState) -> str:
    rows_html = []
    for i, row in enumerate(table.rows):
        cell_tag = "th" if i == 0 else "td"
        cells_html = []
        for cell in _dedupe_row_cells(row):
            span = _cell_grid_span(cell)
            span_attr = f' colspan="{span}"' if span > 1 else ""
            cells_html.append(f"<{cell_tag}{span_attr}>{_cell_html(cell, resolver, state)}</{cell_tag}>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return "<table>\n" + "\n".join(rows_html) + "\n</table>"


def _cell_html(cell, resolver: InlineResolver, state: ConvertState) -> str:
    blocks: list[str] = []
    for item in cell.iter_inner_content():
        if isinstance(item, Paragraph):
            text = _paragraph_inline_html(item, resolver, state)
            if text.strip():
                blocks.append(f"<p>{text}</p>")
        elif isinstance(item, Table):
            blocks.append(_table_to_html(item, resolver, state))
    return "".join(blocks)


def _dedupe_row_cells(row) -> list:
    """python-docx repeats a horizontally-merged (gridSpan) cell once per
    spanned grid column; drop the repeats so its text isn't duplicated."""
    cells = []
    prev_tc = None
    for cell in row.cells:
        if cell._tc is prev_tc:
            continue
        prev_tc = cell._tc
        cells.append(cell)
    return cells


def table_to_markdown(table: Table, resolver: InlineResolver, state: ConvertState) -> list[str]:
    has_nested = any(cell.tables for row in table.rows for cell in row.cells)
    if has_nested:
        return [_table_to_html(table, resolver, state)]

    rows: list[list[str]] = []
    for row in table.rows:
        values = []
        for cell in _dedupe_row_cells(row):
            cell_lines = [paragraph_inline_markdown(p, resolver, state).strip() for p in cell.paragraphs]
            cell_text = "<br>".join(line for line in cell_lines if line)
            values.append(cell_text.replace("|", "\\|"))
            values.extend([""] * (_cell_grid_span(cell) - 1))
        rows.append(values)

    if not rows:
        return []

    col_count = max(len(r) for r in rows)
    normalized = [r + [""] * (col_count - len(r)) for r in rows]

    header = normalized[0]
    sep = ["---"] * col_count
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for data_row in normalized[1:]:
        lines.append("| " + " | ".join(data_row) + " |")

    return lines


# --------------------------------------------------------------------------
# Headers / footers
# --------------------------------------------------------------------------


def _hf_has_content(hf) -> bool:
    for p in hf.paragraphs:
        if (p.text or "").strip():
            return True
        if p._p.xpath(".//w:drawing") or p._p.xpath(".//*[local-name()='pict']"):
            return True
    return bool(hf.tables)


def header_footer_to_markdown(hf, resolver: InlineResolver, state: ConvertState) -> list[str]:
    lines: list[str] = []
    for item in hf.iter_inner_content():
        if isinstance(item, Paragraph):
            paragraph_lines = paragraph_to_markdown(item, resolver, state)
            if paragraph_lines:
                lines.extend(paragraph_lines)
                lines.append("")
        elif isinstance(item, Table):
            table_lines = table_to_markdown(item, resolver, state)
            if table_lines:
                lines.extend(table_lines)
                lines.append("")
    return lines


def _collect_header_footer_blocks(doc, state: ConvertState) -> tuple[list[str], list[str]]:
    header_blocks: list[str] = []
    footer_blocks: list[str] = []
    seen_header_ids: set[int] = set()
    seen_footer_ids: set[int] = set()
    multiple_sections = len(doc.sections) > 1

    for idx, section in enumerate(doc.sections, start=1):
        header = section.header
        if not header.is_linked_to_previous and id(header._element) not in seen_header_ids:
            seen_header_ids.add(id(header._element))
            if _hf_has_content(header):
                if multiple_sections:
                    header_blocks.extend([f"### Section {idx}", ""])
                header_blocks.extend(header_footer_to_markdown(header, _part_resolver(header.part), state))

        footer = section.footer
        if not footer.is_linked_to_previous and id(footer._element) not in seen_footer_ids:
            seen_footer_ids.add(id(footer._element))
            if _hf_has_content(footer):
                if multiple_sections:
                    footer_blocks.extend([f"### Section {idx}", ""])
                footer_blocks.extend(header_footer_to_markdown(footer, _part_resolver(footer.part), state))

    return header_blocks, footer_blocks


# --------------------------------------------------------------------------
# Footnotes / endnotes (python-docx has no read API, so parse the raw parts)
# --------------------------------------------------------------------------


def _load_notes_part(docx_path: Path, part_name: str):
    try:
        with zipfile.ZipFile(docx_path) as z:
            names = set(z.namelist())
            xml_name = f"word/{part_name}.xml"
            if xml_name not in names:
                return None, {}
            root = parse_xml(z.read(xml_name))

            rels: dict[str, str] = {}
            rels_name = f"word/_rels/{part_name}.xml.rels"
            if rels_name in names:
                rels_root = etree.fromstring(z.read(rels_name))
                for rel in rels_root:
                    rid, target = rel.get("Id"), rel.get("Target")
                    if rid and target:
                        rels[rid] = target
            return root, rels
    except (KeyError, zipfile.BadZipFile, etree.XMLSyntaxError):
        return None, {}


def _parse_notes(root, rels: dict[str, str], tag_name: str) -> dict[str, str]:
    notes: dict[str, str] = {}
    if root is None:
        return notes

    resolver = _rels_resolver(rels)
    dummy_state = ConvertState(image_dir=Path("."))

    for note in root.findall(qn(f"w:{tag_name}")):
        if note.get(qn("w:type")) in NOTE_SKIP_TYPES:
            continue
        note_id = note.get(qn("w:id"))
        if note_id is None:
            continue

        parts_text = []
        for p in note.findall(qn("w:p")):
            text = segments_to_markdown(_walk_inline(p, resolver), resolver, dummy_state).strip()
            if text:
                parts_text.append(text)
        notes[note_id] = " ".join(parts_text)

    return notes


# --------------------------------------------------------------------------
# DOCX conversion
# --------------------------------------------------------------------------


def convert_docx(input_path: Path, output_root: Path, base: Optional[str] = None) -> Path:
    doc = Document(str(input_path))
    base = base or sanitize_name(input_path.stem)

    out_dir = output_root / base
    reset_dir(out_dir)

    output_md = out_dir / f"{base}.md"
    image_dir = out_dir / "images"
    ensure_dir(image_dir)

    state = ConvertState(image_dir=image_dir)
    resolver = _part_resolver(doc.part)

    body_lines: list[str] = []
    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            paragraph_lines = paragraph_to_markdown(item, resolver, state)
            if paragraph_lines:
                body_lines.extend(paragraph_lines)
                body_lines.append("")
        elif isinstance(item, Table):
            table_lines = table_to_markdown(item, resolver, state)
            if table_lines:
                body_lines.extend(table_lines)
                body_lines.append("")

    header_blocks, footer_blocks = _collect_header_footer_blocks(doc, state)

    lines: list[str] = [f"# {input_path.stem}", ""]
    if header_blocks:
        lines.extend(["## Header", ""])
        lines.extend(header_blocks)
        lines.extend(["---", ""])

    lines.extend(body_lines)

    if footer_blocks:
        lines.extend(["---", "", "## Footer", ""])
        lines.extend(footer_blocks)

    if state.used_footnotes or state.used_endnotes:
        footnotes_root, footnotes_rels = _load_notes_part(input_path, "footnotes")
        endnotes_root, endnotes_rels = _load_notes_part(input_path, "endnotes")
        footnotes_map = _parse_notes(footnotes_root, footnotes_rels, "footnote")
        endnotes_map = _parse_notes(endnotes_root, endnotes_rels, "endnote")

        lines.extend(["---", ""])
        for fid in state.used_footnotes:
            lines.append(f"[^fn{fid}]: {footnotes_map.get(fid, '')}")
        for eid in state.used_endnotes:
            lines.append(f"[^en{eid}]: {endnotes_map.get(eid, '')}")
        lines.append("")

    if state.image_index == 0 and image_dir.exists() and not any(image_dir.iterdir()):
        image_dir.rmdir()

    write_text(output_md, "\n".join(lines))
    return output_md


# --------------------------------------------------------------------------
# PDF conversion
# --------------------------------------------------------------------------


def convert_pdf(input_path: Path, output_root: Path, base: Optional[str] = None) -> Path:
    base = base or sanitize_name(input_path.stem)

    # pymupdf4llm's internal path sanitizer (spaces/parens -> "_"/"-") builds
    # the on-disk save path from the *sanitized* string while the real folder
    # is created unsanitized, so a space or paren in image_path makes pix.save()
    # fail with "No such file or directory". Route images through a temp dir
    # with no such characters, then move them into images/ ourselves with
    # sequential names, matching the docx/excel image-naming convention.
    with tempfile.TemporaryDirectory(prefix="pdfimg_") as tmp_dir:
        body = pymupdf4llm.to_markdown(
            str(input_path),
            write_images=True,
            image_path=tmp_dir,
            image_format="png",
            dpi=150,
        ).strip()

        # Parsing succeeded: only now create the output folder, so a failed
        # (corrupt/empty) PDF never leaves stray empty folders behind.
        out_dir = output_root / base
        reset_dir(out_dir)
        output_md = out_dir / f"{base}.md"
        image_dir = out_dir / "images"

        tmp_prefix = re.escape(Path(tmp_dir).resolve().as_posix())
        image_index = 0

        def _relink(match: "re.Match[str]") -> str:
            nonlocal image_index
            src_path = Path(tmp_dir) / match.group(1)
            if not src_path.is_file():
                return match.group(0)
            ensure_dir(image_dir)
            image_index += 1
            dest_name = f"img_{image_index:04d}{src_path.suffix or '.png'}"
            shutil.copyfile(src_path, image_dir / dest_name)
            return f"images/{dest_name}"

        body = re.sub(tmp_prefix + r"/([^)\s]+)", _relink, body)

    lines = [f"# {input_path.stem}", "", body, ""]

    write_text(output_md, "\n".join(lines))
    return output_md


# --------------------------------------------------------------------------
# Excel conversion
# --------------------------------------------------------------------------


def excel_sheet_used_range(ws) -> tuple[int, int, int, int]:
    """Return min_row, max_row, min_col, max_col for non-empty cells."""
    min_row: Optional[int] = None
    max_row: Optional[int] = None
    min_col: Optional[int] = None
    max_col: Optional[int] = None

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            if cell.value is None:
                continue
            r, c = cell.row, cell.column
            min_row = r if min_row is None else min(min_row, r)
            max_row = r if max_row is None else max(max_row, r)
            min_col = c if min_col is None else min(min_col, c)
            max_col = c if max_col is None else max(max_col, c)

    if min_row is None:
        return 1, 1, 1, 1
    return min_row, max_row or 1, min_col or 1, max_col or 1


def format_cell(value) -> str:
    if value is None:
        return ""
    text = str(value).replace("\n", " ").strip()
    text = text.replace("|", "\\|")
    return text


def extract_excel_images(ws, images_dir: Path, prefix: str, image_index: int) -> tuple[list[str], int]:
    lines: list[str] = []
    images = getattr(ws, "_images", [])

    for img in images:
        ext = ".png"
        fmt = getattr(img, "format", None)
        if isinstance(fmt, str) and fmt:
            ext = f".{fmt.lower()}"

        image_index += 1
        filename = f"{prefix}_img_{image_index:04d}{ext}"
        target = images_dir / filename

        try:
            data = img._data()
            blob = data() if callable(data) else data
            if blob is None:
                continue
            target.write_bytes(blob)
        except Exception:
            # Skip non-exportable images, keep conversion resilient.
            continue

        anchor = getattr(img, "anchor", None)
        position = ""
        if hasattr(anchor, "_from"):
            row = anchor._from.row + 1
            col = anchor._from.col + 1
            position = f" (cell R{row}C{col})"

        lines.append(f"![{ws.title}{position}](images/{filename})")

    return lines, image_index


def convert_excel(input_path: Path, output_root: Path, base: Optional[str] = None) -> Path:
    wb = load_workbook(filename=str(input_path), data_only=True)
    base = base or sanitize_name(input_path.stem)

    out_dir = output_root / base
    reset_dir(out_dir)

    images_dir = out_dir / "images"
    ensure_dir(images_dir)

    image_index = 0

    for ws in wb.worksheets:
        sheet_name = sanitize_name(ws.title)
        md_path = out_dir / f"{sheet_name}.md"

        min_row, max_row, min_col, max_col = excel_sheet_used_range(ws)
        width = max_col - min_col + 1

        lines: list[str] = [f"# {ws.title}", ""]

        headers = [
            format_cell(ws.cell(row=min_row, column=min_col + i).value) or f"Column {i + 1}"
            for i in range(width)
        ]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * width) + " |")

        for r in range(min_row + 1, max_row + 1):
            row_data = [format_cell(ws.cell(row=r, column=min_col + i).value) for i in range(width)]
            lines.append("| " + " | ".join(row_data) + " |")

        safe_prefix = sheet_name.replace(" ", "_")
        image_lines, image_index = extract_excel_images(ws, images_dir, safe_prefix, image_index)
        if image_lines:
            lines.append("")
            lines.append("## Images")
            lines.append("")
            lines.extend(image_lines)

        write_text(md_path, "\n".join(lines))

    if image_index == 0 and images_dir.exists() and not any(images_dir.iterdir()):
        images_dir.rmdir()

    return out_dir


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------


def iter_input_files(path: Path) -> Iterable[Path]:
    if path.is_file():
        yield path
        return

    for ext in ("*.docx", "*.xlsx", "*.xlsm", "*.pdf"):
        yield from path.rglob(ext)


def convert_file(input_path: Path, output_dir: Path, base: Optional[str] = None) -> Path:
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        return convert_docx(input_path, output_dir, base)
    if suffix in {".xlsx", ".xlsm"}:
        return convert_excel(input_path, output_dir, base)
    if suffix == ".pdf":
        return convert_pdf(input_path, output_dir, base)
    raise ValueError(f"Unsupported file type: {input_path}")


def _resolve_output_bases(file_paths: list[Path]) -> dict[Path, str]:
    """Assign each input file a unique output base name. Different files whose
    names collide after sanitize_name() (e.g. "Report?.docx" and "Report*.docx"
    both -> "Report_") would otherwise share one output folder and silently
    clobber each other's output."""
    seen: dict[str, int] = {}
    resolved: dict[Path, str] = {}
    for file_path in file_paths:
        base = sanitize_name(file_path.stem)
        count = seen.get(base, 0) + 1
        seen[base] = count
        resolved[file_path] = base if count == 1 else f"{base} ({count})"
    return resolved


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Convert DOCX, Excel, or PDF files to Markdown with extracted images."
    )
    parser.add_argument("input", type=Path, help="Input file or folder")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output root folder (default: same folder as input)",
    )
    args = parser.parse_args()

    input_path = args.input.expanduser().resolve()
    if not input_path.exists():
        print(f"Input not found: {input_path}")
        return 1

    output_root = (
        args.output.expanduser().resolve()
        if args.output
        else (input_path.parent if input_path.is_file() else input_path)
    )
    try:
        ensure_dir(output_root)
    except OSError as exc:
        print(f"Cannot create output folder {output_root}: {exc}")
        return 1

    file_paths: list[Path] = []
    seen: set[Path] = set()
    for file_path in iter_input_files(input_path):
        if file_path.name.startswith("~$") or file_path in seen:
            continue
        seen.add(file_path)
        file_paths.append(file_path)

    output_bases = _resolve_output_bases(file_paths)

    converted: list[Path] = []
    failed: list[tuple[Path, Exception]] = []
    for file_path in file_paths:
        try:
            converted.append(convert_file(file_path, output_root, output_bases[file_path]))
        except Exception as exc:  # keep processing the rest of the batch
            failed.append((file_path, exc))

    if not converted and not failed:
        print("No supported files found (.docx, .xlsx, .xlsm, .pdf).")
        return 0

    if converted:
        print("Converted outputs:")
        for path in converted:
            print(f"- {path}")

    if failed:
        print("Failed to convert:")
        for path, exc in failed:
            print(f"- {path}: {exc}")

    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
